"""
Document Generator Module
Generates tailored resume in PDF and DOCX formats.
"""

import io
from pathlib import Path
from typing import Union, Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER


def generate_docx(resume_text: str, output_path: Union[str, Path] = None) -> Union[Document, bytes]:
    """
    Generate DOCX document from resume text.

    Args:
        resume_text: Tailored resume text
        output_path: Optional path to save document

    Returns:
        Document object or bytes if no output_path
    """
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Parse resume text and add to document
    lines = resume_text.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            # Add spacing for empty lines
            doc.add_paragraph()
            continue

        # Detect headers (all caps or specific keywords)
        is_header = (
            line.isupper() or
            any(keyword in line.upper() for keyword in [
                'EXPERIENCE', 'EDUCATION', 'SKILLS', 'PROJECTS',
                'CERTIFICATIONS', 'SUMMARY', 'OBJECTIVE'
            ])
        )

        # Detect name (first line, typically)
        is_name = len(doc.paragraphs) == 0 and len(line) < 50

        # Detect bullet points
        is_bullet = line.startswith('•') or line.startswith('-') or line.startswith('*')

        if is_name:
            # Format name
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(18)
            run.font.bold = True

        elif is_header:
            # Format section header
            p = doc.add_paragraph(line.upper())
            run = p.runs[0]
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue

        elif is_bullet:
            # Format bullet point
            # Remove existing bullet markers
            clean_line = line.lstrip('•-* ').strip()
            p = doc.add_paragraph(clean_line, style='List Bullet')
            run = p.runs[0]
            run.font.size = Pt(10)

        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(10)

    # Save or return bytes
    if output_path:
        doc.save(output_path)
        return doc
    else:
        # Return as bytes
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()


def generate_pdf(resume_text: str, output_path: Union[str, Path] = None) -> bytes:
    """
    Generate PDF document from resume text.

    Args:
        resume_text: Tailored resume text
        output_path: Optional path to save document

    Returns:
        PDF bytes
    """
    # Create buffer or file
    if output_path:
        buffer = str(output_path)
    else:
        buffer = io.BytesIO()

    # Create PDF document
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch
    )

    # Styles
    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        name='CustomName',
        parent=styles['Heading1'],
        fontSize=18,
        textColor='black',
        alignment=TA_CENTER,
        spaceAfter=12
    ))

    styles.add(ParagraphStyle(
        name='CustomHeader',
        parent=styles['Heading2'],
        fontSize=12,
        textColor='darkblue',
        bold=True,
        spaceAfter=6,
        spaceBefore=12
    ))

    styles.add(ParagraphStyle(
        name='CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6
    ))

    styles.add(ParagraphStyle(
        name='CustomBullet',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        bulletIndent=10,
        spaceAfter=4
    ))

    # Build document content
    story = []
    lines = resume_text.split('\n')

    for i, line in enumerate(lines):
        line = line.strip()

        if not line:
            story.append(Spacer(1, 0.1 * inch))
            continue

        # Detect different sections
        is_header = (
            line.isupper() or
            any(keyword in line.upper() for keyword in [
                'EXPERIENCE', 'EDUCATION', 'SKILLS', 'PROJECTS',
                'CERTIFICATIONS', 'SUMMARY', 'OBJECTIVE'
            ])
        )

        is_name = i == 0 and len(line) < 50

        is_bullet = line.startswith('•') or line.startswith('-') or line.startswith('*')

        if is_name:
            # Name at top
            p = Paragraph(line, styles['CustomName'])
            story.append(p)

        elif is_header:
            # Section header
            p = Paragraph(line.upper(), styles['CustomHeader'])
            story.append(p)

        elif is_bullet:
            # Bullet point
            clean_line = line.lstrip('•-* ').strip()
            bullet_text = f"• {clean_line}"
            p = Paragraph(bullet_text, styles['CustomBullet'])
            story.append(p)

        else:
            # Regular text
            p = Paragraph(line, styles['CustomBody'])
            story.append(p)

    # Build PDF
    doc.build(story)

    # Return bytes if no output path
    if not output_path:
        buffer.seek(0)
        return buffer.getvalue()

    return None


def generate_resume_document(
    resume_text: str,
    format: str = 'both',
    output_dir: Union[str, Path] = None,
    filename_base: str = 'tailored_resume'
) -> Dict:
    """
    Generate resume document in specified format(s).

    Args:
        resume_text: Tailored resume text
        format: 'pdf', 'docx', or 'both'
        output_dir: Directory to save files (if None, returns bytes)
        filename_base: Base filename without extension

    Returns:
        Dictionary with file paths or bytes
    """
    result = {
        'success': True,
        'formats': []
    }

    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True, parents=True)

    try:
        if format in ['docx', 'both']:
            if output_dir:
                docx_path = output_dir / f"{filename_base}.docx"
                generate_docx(resume_text, docx_path)
                result['docx_path'] = str(docx_path)
            else:
                result['docx_bytes'] = generate_docx(resume_text)
            result['formats'].append('docx')

        if format in ['pdf', 'both']:
            if output_dir:
                pdf_path = output_dir / f"{filename_base}.pdf"
                generate_pdf(resume_text, pdf_path)
                result['pdf_path'] = str(pdf_path)
            else:
                result['pdf_bytes'] = generate_pdf(resume_text)
            result['formats'].append('pdf')

    except Exception as e:
        result['success'] = False
        result['error'] = str(e)

    return result


def create_comparison_document(
    original_text: str,
    tailored_text: str,
    summary: str,
    scores: Dict,
    output_path: Union[str, Path]
) -> bool:
    """
    Create a comparison document showing original vs tailored resume.

    Args:
        original_text: Original resume
        tailored_text: Tailored resume
        summary: Tailoring summary
        scores: Dictionary with initial_score and final_score
        output_path: Path to save document

    Returns:
        True if successful
    """
    try:
        doc = Document()

        # Title
        doc.add_heading('Resume Tailoring Report', 0)

        # Summary section
        doc.add_heading('Tailoring Summary', 1)
        doc.add_paragraph(summary)

        # Scores
        doc.add_heading('Match Scores', 1)
        doc.add_paragraph(f"Initial Score: {scores.get('initial_score', 0):.1f}%")
        doc.add_paragraph(f"Final Score: {scores.get('final_score', 0):.1f}%")
        doc.add_paragraph(f"Improvement: +{scores.get('improvement', 0):.1f}%")

        # Original resume
        doc.add_page_break()
        doc.add_heading('Original Resume', 1)
        doc.add_paragraph(original_text)

        # Tailored resume
        doc.add_page_break()
        doc.add_heading('Tailored Resume', 1)
        doc.add_paragraph(tailored_text)

        doc.save(output_path)
        return True

    except Exception as e:
        print(f"Error creating comparison document: {e}")
        return False
