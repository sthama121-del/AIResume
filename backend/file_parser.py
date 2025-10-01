"""
File Parser Module
Handles extraction of text from various file formats (PDF, DOCX, TXT).
"""

import io
from pathlib import Path
from typing import Union, Dict
import PyPDF2
import pdfplumber
import fitz  # PyMuPDF
from docx import Document


def extract_text_from_pdf(file_path: Union[str, Path]) -> str:
    """
    Extract text from PDF file using multiple methods for better accuracy.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text as string
    """
    text = ""

    # Method 1: Try pdfplumber first (best for formatted resumes)
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber extraction failed: {e}")

    # Method 2: If pdfplumber fails or gets minimal text, try PyMuPDF
    if len(text.strip()) < 100:
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
        except Exception as e:
            print(f"PyMuPDF extraction failed: {e}")

    # Method 3: Fallback to PyPDF2
    if len(text.strip()) < 100:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")

    return text.strip()


def extract_text_from_docx(file_path: Union[str, Path]) -> str:
    """
    Extract text from DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text as string
    """
    try:
        doc = Document(file_path)
        text = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        return "\n".join(text)
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {e}")


def extract_text_from_txt(file_path: Union[str, Path]) -> str:
    """
    Extract text from TXT file.

    Args:
        file_path: Path to the TXT file

    Returns:
        Extracted text as string
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        with open(file_path, 'r', encoding='latin-1') as file:
            return file.read()


def extract_text_from_uploaded_file(uploaded_file) -> str:
    """
    Extract text from Streamlit uploaded file object.

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        Extracted text as string
    """
    file_extension = uploaded_file.name.split('.')[-1].lower()

    if file_extension == 'pdf':
        # Save to temporary file for PDF processing
        temp_path = Path(f"/tmp/{uploaded_file.name}")
        with open(temp_path, 'wb') as f:
            f.write(uploaded_file.getbuffer())
        text = extract_text_from_pdf(temp_path)
        temp_path.unlink()  # Delete temp file
        return text

    elif file_extension == 'docx':
        # Process DOCX from bytes
        doc = Document(io.BytesIO(uploaded_file.getvalue()))
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        return "\n".join(text)

    elif file_extension == 'txt':
        return uploaded_file.getvalue().decode('utf-8')

    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


def extract_text_from_file(file_path: Union[str, Path]) -> str:
    """
    Main function to extract text from any supported file format.

    Args:
        file_path: Path to the file

    Returns:
        Extracted text as string
    """
    file_path = Path(file_path)
    file_extension = file_path.suffix.lower()

    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


def parse_resume_structure(resume_text: str) -> Dict:
    """
    Parse resume text to identify structure (projects, dates, bullets).

    Args:
        resume_text: Raw resume text

    Returns:
        Dictionary with parsed resume structure
    """
    lines = resume_text.split('\n')
    structure = {
        'raw_text': resume_text,
        'lines': lines,
        'projects': [],
        'experience_section': []
    }

    # Simple heuristic to find experience/project sections
    # This is a basic implementation - can be enhanced with NLP
    experience_keywords = ['experience', 'work history', 'employment', 'projects', 'professional experience']

    in_experience_section = False
    current_project = None

    for i, line in enumerate(lines):
        line_lower = line.lower().strip()

        # Detect experience section start
        if any(keyword in line_lower for keyword in experience_keywords):
            in_experience_section = True
            continue

        if in_experience_section:
            structure['experience_section'].append(line)

    return structure
